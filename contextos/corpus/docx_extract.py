"""docx bytes -> 正文文本 + 嵌入图 bytes。

正文走 python-docx(段落 + 表格); 嵌入图走 zipfile(word/media/), 不依赖 python-docx
的图片关系解析(更稳)。统一 bytes 接口便于测试与缓存。
"""
from __future__ import annotations

import io
import zipfile

from docx import Document


def extract_text(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)
    return "\n".join(parts)


def extract_images(data: bytes) -> list[tuple[str, bytes]]:
    """返回 [(文件名, 图字节), ...], 按 zip 内名字排序(稳定顺序)。"""
    out: list[tuple[str, bytes]] = []
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        for name in sorted(z.namelist()):
            if name.startswith("word/media/"):
                out.append((name.split("/")[-1], z.read(name)))
    return out
