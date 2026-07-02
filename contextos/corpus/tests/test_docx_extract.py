import base64

from docx import Document

# 1x1 RGB PNG(合法最小 PNG, python-docx add_picture 可接受)
_PNG_1x1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAIAAACQd1PeAAAADElEQVR4nGP4z8AAAAMBAQDJ/pLvAAAAAElFTkSuQmCC"
)


def _make_docx_bytes(tmp_path, paragraphs, table_rows=None, with_image=False):
    d = Document()
    for p in paragraphs:
        d.add_paragraph(p)
    if table_rows:
        t = d.add_table(rows=0, cols=len(table_rows[0]))
        for row in table_rows:
            cells = t.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = val
    if with_image:
        import io
        d.add_picture(io.BytesIO(_PNG_1x1))
    p = tmp_path / "x.docx"
    d.save(str(p))
    return p.read_bytes()


def test_extract_text_paragraphs(tmp_path):
    from contextos.corpus.docx_extract import extract_text
    data = _make_docx_bytes(tmp_path, ["Province tax config", "Click RULE_TYPE_TAX"])
    txt = extract_text(data)
    assert "Province tax config" in txt
    assert "RULE_TYPE_TAX" in txt


def test_extract_text_includes_tables(tmp_path):
    from contextos.corpus.docx_extract import extract_text
    data = _make_docx_bytes(
        tmp_path, ["header"], table_rows=[["TableName", "PK"], ["CONF_PROVINCE_TAX", "PROVINCE"]]
    )
    txt = extract_text(data)
    assert "CONF_PROVINCE_TAX" in txt
    assert "PROVINCE" in txt


def test_extract_images_returns_embedded(tmp_path):
    from contextos.corpus.docx_extract import extract_images
    data = _make_docx_bytes(tmp_path, ["with image"], with_image=True)
    imgs = extract_images(data)
    assert len(imgs) == 1
    name, blob = imgs[0]
    assert name.startswith("image")
    assert blob[:4] == b"\x89PNG"


def test_extract_images_none_when_no_media(tmp_path):
    from contextos.corpus.docx_extract import extract_images
    data = _make_docx_bytes(tmp_path, ["no image"])
    assert extract_images(data) == []
