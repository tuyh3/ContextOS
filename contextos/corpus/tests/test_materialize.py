import base64

from docx import Document

# VALID 1x1 PNG (per CRITICAL FIX instruction)
_PNG_1x1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAIAAACQd1PeAAAADElEQVR4nGP4z8AAAAMBAQDJ/pLvAAAAAElFTkSuQmCC"
)


def _docx_bytes(tmp_path, paragraphs, with_image=False):
    d = Document()
    for p in paragraphs:
        d.add_paragraph(p)
    if with_image:
        import io
        d.add_picture(io.BytesIO(_PNG_1x1))
    p = tmp_path / "x.docx"
    d.save(str(p))
    return p.read_bytes()


def test_content_hash_only_content(tmp_path):
    """同 content -> 同 hash(不掺路径/时间)。"""
    from contextos.corpus.materialize import content_hash
    assert content_hash(b"abc") == content_hash(b"abc")
    assert content_hash(b"abc") != content_hash(b"abd")


def test_sidecar_md_passthrough():
    from contextos.corpus.materialize import build_sidecar_text
    from contextos.corpus.ocr.fake import FakeOcr
    out = build_sidecar_text(b"# Title\nbody", "md", FakeOcr())
    assert "# Title" in out and "body" in out


def test_sidecar_png_is_ocr():
    from contextos.corpus.materialize import build_sidecar_text
    from contextos.corpus.ocr.fake import FakeOcr
    out = build_sidecar_text(b"pngbytes", "png", FakeOcr(default_text="CONF_PROVINCE_TAX"))
    assert "CONF_PROVINCE_TAX" in out
    assert "[image" in out  # OCR 来源标记


def test_sidecar_docx_text_plus_ocr(tmp_path):
    from contextos.corpus.materialize import build_sidecar_text
    from contextos.corpus.ocr.fake import FakeOcr
    data = _docx_bytes(tmp_path, ["Province tax config"], with_image=True)
    out = build_sidecar_text(data, "docx", FakeOcr(default_text="RULE_TYPE_TAX"))
    assert "Province tax config" in out      # 正文
    assert "RULE_TYPE_TAX" in out            # 截图 OCR
    assert "[image 1 OCR" in out             # 来源标记


def _make_corpus(tmp_path):
    """造一个 dir 源: 1 md + 1 docx(带图) + 1 xlsx(应被 format filter 跳过 —— 默认 formats 不含 xlsx)。"""
    src_dir = tmp_path / "src"
    (src_dir / "doc").mkdir(parents=True)
    (src_dir / "doc" / "note.md").write_text("# Note\nDynamicChargingSVImpl", encoding="utf-8")
    (src_dir / "doc" / "cfg.docx").write_bytes(_docx_bytes(tmp_path, ["table config"], with_image=True))
    (src_dir / "doc" / "fpa.xlsx").write_bytes(b"PK\x03\x04 fake xlsx")
    return src_dir


def _store(tmp_path):
    from contextos.storage.db import make_engine
    from contextos.corpus.record_store import RecordStore
    return RecordStore(make_engine(f"sqlite:///{tmp_path / 'rec.db'}"))


def test_materialize_corpus_writes_sidecars_and_skips_xlsx(tmp_path):
    from contextos.profile.schema import SourceConfig
    from contextos.corpus.materialize import materialize_corpus
    from contextos.corpus.ocr.fake import FakeOcr
    src_dir = _make_corpus(tmp_path)
    out_dir = tmp_path / "materialized"
    src = SourceConfig(type="dir", location=str(src_dir),
                       glob=["**/*.md", "**/*.docx", "**/*.xlsx"])
    stats = materialize_corpus(
        sources=[src], materialized_dir=out_dir, store=_store(tmp_path),
        ocr=FakeOcr(default_text="OCRTABLE"), backend_name="fake",
    )
    # md + docx 物化; xlsx 被 format filter 跳过(不在默认 formats; xlsx 物化路径留 v2)
    assert (out_dir / "doc" / "note.md").exists()
    assert (out_dir / "doc" / "cfg.docx.md").exists()
    assert not (out_dir / "doc" / "fpa.xlsx.md").exists()
    assert "OCRTABLE" in (out_dir / "doc" / "cfg.docx.md").read_text(encoding="utf-8")
    assert stats["materialized"] == 2
    assert stats["skipped"] == 0


def test_materialize_corpus_skips_unchanged_on_second_run(tmp_path):
    from contextos.profile.schema import SourceConfig
    from contextos.corpus.materialize import materialize_corpus
    from contextos.corpus.ocr.fake import FakeOcr
    src_dir = _make_corpus(tmp_path)
    out_dir = tmp_path / "materialized"
    store = _store(tmp_path)
    src = SourceConfig(type="dir", location=str(src_dir), glob=["**/*.md", "**/*.docx"])
    materialize_corpus(sources=[src], materialized_dir=out_dir, store=store,
                       ocr=FakeOcr(), backend_name="fake")
    stats2 = materialize_corpus(sources=[src], materialized_dir=out_dir, store=store,
                                ocr=FakeOcr(), backend_name="fake")
    assert stats2["materialized"] == 0
    assert stats2["skipped"] == 2


def test_materialize_corpus_full_cleanup_removes_deleted(tmp_path):
    from contextos.profile.schema import SourceConfig
    from contextos.corpus.materialize import materialize_corpus
    from contextos.corpus.ocr.fake import FakeOcr
    src_dir = _make_corpus(tmp_path)
    out_dir = tmp_path / "materialized"
    store = _store(tmp_path)
    src = SourceConfig(type="dir", location=str(src_dir), glob=["**/*.md", "**/*.docx"])
    materialize_corpus(sources=[src], materialized_dir=out_dir, store=store,
                       ocr=FakeOcr(), backend_name="fake")
    # 删源里的 md, 再跑 -> sidecar + record 应被清
    (src_dir / "doc" / "note.md").unlink()
    materialize_corpus(sources=[src], materialized_dir=out_dir, store=store,
                       ocr=FakeOcr(), backend_name="fake")
    assert not (out_dir / "doc" / "note.md").exists()
    assert "doc/note.md" not in store.all_doc_ids()


def test_materialize_corpus_rematerializes_when_sidecar_copy_missing(tmp_path):
    """store/磁盘失同步: hash 命中但 sidecar 副本被删 -> 必须重新物化, 不能靠 record
    残留静默跳过(否则副本丢失后语料永久缺该文件)。守 materialize.py 的
    '(materialized_dir / rec.sidecar_path).exists()' 分支。"""
    from contextos.profile.schema import SourceConfig
    from contextos.corpus.materialize import materialize_corpus
    from contextos.corpus.ocr.fake import FakeOcr
    src_dir = _make_corpus(tmp_path)
    out_dir = tmp_path / "materialized"
    store = _store(tmp_path)
    src = SourceConfig(type="dir", location=str(src_dir), glob=["**/*.md"])
    materialize_corpus(sources=[src], materialized_dir=out_dir, store=store,
                       ocr=FakeOcr(), backend_name="fake")
    sidecar = out_dir / "doc" / "note.md"
    assert sidecar.exists()
    sidecar.unlink()                          # 磁盘副本丢失, record 仍在(失同步)
    stats = materialize_corpus(sources=[src], materialized_dir=out_dir, store=store,
                               ocr=FakeOcr(), backend_name="fake")
    assert sidecar.exists()                   # 重新物化补回, 不是被跳过
    assert stats["materialized"] == 1
    assert stats["skipped"] == 0


def test_materialize_corpus_failsafe_one_bad_file(tmp_path):
    """单文档物化失败(坏 docx)不挂全局, 好文档照常物化。"""
    from contextos.profile.schema import SourceConfig
    from contextos.corpus.materialize import materialize_corpus
    from contextos.corpus.ocr.fake import FakeOcr
    src_dir = tmp_path / "src"
    (src_dir / "doc").mkdir(parents=True)
    (src_dir / "doc" / "good.md").write_text("ok", encoding="utf-8")
    (src_dir / "doc" / "bad.docx").write_bytes(b"not a real docx zip")
    out_dir = tmp_path / "materialized"
    src = SourceConfig(type="dir", location=str(src_dir), glob=["**/*.md", "**/*.docx"])
    stats = materialize_corpus(sources=[src], materialized_dir=out_dir, store=_store(tmp_path),
                               ocr=FakeOcr(), backend_name="fake")
    assert (out_dir / "doc" / "good.md").exists()
    assert stats["failed"] == 1
    assert stats["materialized"] == 1
