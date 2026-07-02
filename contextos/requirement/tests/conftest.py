from __future__ import annotations

from pathlib import Path

import pytest


@pytest.fixture
def make_docx(tmp_path):
    """造一个含段落 + 表格的 .docx,返回路径。"""
    def _make(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> Path:
        from docx import Document

        doc = Document()
        for p in paragraphs:
            doc.add_paragraph(p)
        if table_rows:
            rows = len(table_rows)
            cols = max(len(r) for r in table_rows)
            table = doc.add_table(rows=rows, cols=cols)
            for i, r in enumerate(table_rows):
                for j, cell_text in enumerate(r):
                    table.rows[i].cells[j].text = cell_text
        out = tmp_path / "sample.docx"
        doc.save(out)
        return out

    return _make
