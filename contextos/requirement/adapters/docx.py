"""docx 源适配器:python-docx 段落 + 表格行列线性化 + tracked 插入抽取。

图片 OCR + 全保真 comments 留 v2(design 02 §1.1)。表格线性化为
"cell | cell | cell"(design 明确表格质量直接影响拆解)。
"""
from __future__ import annotations

from contextos.requirement.adapters.base import AdapterResult, parse_failure, register


def _table_lines(table) -> list[str]:
    lines = []
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        # 去重连续重复 cell(合并单元格 python-docx 会重复同一 cell.text)
        deduped: list[str] = []
        for c in cells:
            if not deduped or deduped[-1] != c:
                deduped.append(c)
        line = " | ".join(deduped).strip()
        if line.strip(" |"):
            lines.append(line)
    return lines


def _tracked_insertions(doc) -> list[str]:
    """抽 w:ins 内的 w:t 文本(track-changes 插入)。"""
    from docx.oxml.ns import qn

    out: list[str] = []
    body = doc.element.body
    for ins in body.iter(qn("w:ins")):
        texts = [t.text for t in ins.iter(qn("w:t")) if t.text]
        joined = "".join(texts).strip()
        if joined:
            out.append(joined)
    return out


def parse_docx(raw_input: str) -> AdapterResult:
    try:
        from docx import Document
    except ImportError:  # pragma: no cover - 依赖已声明
        return parse_failure("python-docx 未安装")

    try:
        doc = Document(raw_input)
    except FileNotFoundError:
        return parse_failure(f"文件不存在: {raw_input}")
    except Exception as e:  # 损坏 / 加密 / 非 docx
        return parse_failure(f"{type(e).__name__}: {e}")

    parts: list[str] = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            parts.append(txt)
    for table in doc.tables:
        parts.extend(_table_lines(table))

    insertions = _tracked_insertions(doc)
    if insertions:
        parts.append("[修订/批注插入]")
        parts.extend(insertions)

    raw_text = "\n".join(parts).strip()
    if not raw_text:
        return parse_failure("空文档(无段落/表格/批注文本)")
    return AdapterResult(raw_text=raw_text)


register("docx", parse_docx)
